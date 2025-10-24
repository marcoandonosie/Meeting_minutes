from docx import Document

def create_doc(meeting_data):
    """
    meeting_data = [date, place, time, title, attendees, context, discussion, next_steps]
    Populates a pre-formatted Word template with these values and saves a new file.
    """

    # Load your Word template
    template_path = "core/EXAMPLE FORMAT MoM Danantara.docx"
    doc = Document(template_path)

    # Define your placeholders and corresponding values
    placeholders = {
        "{{DATE}}": meeting_data[0],
        "{{PLACE}}": meeting_data[1],
        "{{TIME}}": meeting_data[2],
        "{{TITLE}}": meeting_data[3],
        "{{ATTENDEES}}": meeting_data[4],
        "{{CONTEXT}}": meeting_data[5],
        "{{DISCUSSION}}": meeting_data[6],
        "{{NEXTSTEPS}}": meeting_data[7],
    }

    # Replace text in all paragraphs
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Replace text in tables (if the template has tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save output file
    output_path = "Generated_MoM.docx"
    doc.save(output_path)
    print(f"Document generated successfully: {output_path}")

    return output_path